"""
English Learning Automation - Phiên bản Fix Lỗi RSS BBC + Gửi Email
"""
import smtplib
import os
import time
import feedparser
import requests
import google.generativeai as genai
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
from docx import Document
from dotenv import load_dotenv
from bs4 import BeautifulSoup # Thư viện quan trọng để fix lỗi BBC

# --- CẤU HÌNH ---
RSS_FEED_URL = "http://feeds.bbci.co.uk/learningenglish/english/features/6-minute-english/rss"
TEMP_AUDIO_FILE = "temp_podcast.mp3"
MODEL_NAME = "gemini-flash-latest" # Dùng bản Flash cho nhanh và ổn định

class PodcastLearningAutomation:
    def __init__(self):
        self.rss_url = RSS_FEED_URL
        self.setup_env()
        self.setup_gemini()
        
    def setup_env(self):
        """Tải biến môi trường"""
        load_dotenv()
        # Lưu ý: Đảm bảo tên biến khớp với GitHub Secrets của bạn
        self.email_sender = os.getenv("EMAIL_SENDER") or os.getenv("EMAIL_USER")
        self.email_password = os.getenv("EMAIL_PASSWORD") or os.getenv("EMAIL_PASS")
        self.email_receiver = os.getenv("EMAIL_RECEIVER")
        self.api_key = os.getenv('GOOGLE_API_KEY')

        if not self.api_key:
            print("⚠️ Cảnh báo: Thiếu API Key")

    def setup_gemini(self):
        try:
            genai.configure(api_key=self.api_key)
            self.model = genai.GenerativeModel(MODEL_NAME)
        except Exception as e:
            print(f"Lỗi cấu hình Gemini: {e}")

    # --- PHẦN SỬA LỖI (FIX): QUÉT WEB TÌM LINK MP3 ---
    def get_audio_from_webpage(self, page_url):
        """Nếu RSS không có link tải, dùng hàm này để tìm nút Download trên web"""
        try:
            print(f"  🔍 Đang quét trang web tìm link ẩn: {page_url}")
            headers = {'User-Agent': 'Mozilla/5.0'}
            response = requests.get(page_url, headers=headers, timeout=10)
            soup = BeautifulSoup(response.content, 'html.parser')
            
            # Tìm tất cả thẻ <a> có đuôi .mp3
            for a_tag in soup.find_all('a', href=True):
                href = a_tag['href']
                if href.strip().lower().endswith('.mp3'):
                    if href.startswith('/'):
                        return "https://www.bbc.co.uk" + href
                    return href
            return None
        except Exception as e:
            print(f"  ⚠️ Lỗi quét web: {e}")
            return None

    def fetch_latest_episode(self):
        """Logic tải Podcast thông minh hơn"""
        try:
            print(f"\n📡 Đang tải RSS feed...")
            feed = feedparser.parse(self.rss_url)
            if not feed.entries: raise Exception("RSS Trống")
            
            latest = feed.entries[0]
            title = latest.title
            pub_date = latest.get('published', 'Unknown')
            audio_url = None

            # 1. Tìm trong Enclosures (Chuẩn cũ)
            if hasattr(latest, 'enclosures'):
                for enc in latest.enclosures:
                    if enc.get('href', '').endswith('.mp3'):
                        audio_url = enc.get('href'); break
            
            # 2. Tìm trong Media Content (Chuẩn BBC cũ)
            if not audio_url and hasattr(latest, 'media_content'):
                for media in latest.media_content:
                    if media.get('url', '').endswith('.mp3'):
                        audio_url = media.get('url'); break

            # 3. KÍCH HOẠT QUÉT WEB (Giải pháp cho lỗi hiện tại)
            if not audio_url:
                print("  ⚠️ Không thấy link trong RSS, kích hoạt chế độ Web Scraping...")
                audio_url = self.get_audio_from_webpage(latest.link)

            if not audio_url: raise Exception("Không tìm thấy file MP3 bằng mọi cách")

            print(f"✓ Tìm thấy tập: {title}")
            return {'title': title, 'pub_date': pub_date, 'audio_url': audio_url}
        except Exception as e:
            raise Exception(f"Lỗi lấy dữ liệu: {e}")

    def download_audio(self, audio_url):
        print(f"⬇️  Đang tải file MP3...")
        headers = {'User-Agent': 'Mozilla/5.0'} # Giả lập trình duyệt để không bị chặn
        r = requests.get(audio_url, headers=headers, stream=True)
        with open(TEMP_AUDIO_FILE, 'wb') as f:
            for chunk in r.iter_content(chunk_size=8192):
                f.write(chunk)
        return TEMP_AUDIO_FILE

    def process_with_gemini(self, file_path):
        print(f"☁️  Upload lên Gemini & Phân tích...")
        audio_file = genai.upload_file(file_path)
        
        while audio_file.state.name == "PROCESSING":
            time.sleep(5)
            audio_file = genai.get_file(audio_file.name)
            
        if audio_file.state.name == "FAILED": raise Exception("Gemini xử lý thất bại")

        prompt = """
        Bạn là giáo viên tiếng Anh cho người Việt. Hãy phân tích file âm thanh này.
        
        PHẦN 1: PHÂN TÍCH (Analysis)
        1. TỪ VỰNG (5 từ B2-C1):
           - Từ vựng & Loại từ
           - Định nghĩa (Tiếng Việt)
           - Ví dụ & Dịch nghĩa
        2. NGỮ PHÁP (2 cấu trúc):
           - Cấu trúc & Cách dùng (Tiếng Việt)
           - Ví dụ
           
        PHẦN 2: TRANSCRIPT (Bản chép lời đầy đủ)
        
        Định dạng đầu ra rõ ràng để đưa vào file Word.
        """
        response = self.model.generate_content([audio_file, prompt])
        audio_file.delete()
        return response.text

    def create_word_doc(self, info, content):
        print(f"📄 Đang tạo file Word...")
        doc = Document()
        doc.add_heading(info['title'], 0)
        doc.add_paragraph(f"Ngày phát hành: {info['pub_date']}")
        
        # Xử lý nội dung Gemini trả về để đưa vào Word
        for line in content.split('\n'):
            if line.strip():
                if line.startswith('#'):
                    doc.add_heading(line.replace('#', '').strip(), level=2)
                else:
                    doc.add_paragraph(line.strip())
        
        clean_title = "".join([c for c in info['title'] if c.isalnum() or c==' ']).strip().replace(' ', '_')
        filename = f"English_Lesson_{clean_title}.docx"
        doc.save(filename)
        return filename

    def send_email(self, attachment_path, subject):
        if not self.email_sender or not self.email_password:
            print("⚠️ Bỏ qua gửi mail vì thiếu thông tin đăng nhập.")
            return

        print(f"📧 Đang gửi email tới {self.email_receiver}...")
        msg = MIMEMultipart()
        msg['From'] = self.email_sender
        msg['To'] = self.email_receiver
        msg['Subject'] = f"[English Daily] {subject}"
        
        msg.attach(MIMEText("Chào bạn,\n\nĐây là bài học hôm nay. Chúc bạn học vui vẻ!", 'plain'))

        with open(attachment_path, "rb") as f:
            part = MIMEBase('application', 'octet-stream')
            part.set_payload(f.read())
            encoders.encode_base64(part)
            part.add_header('Content-Disposition', f"attachment; filename= {os.path.basename(attachment_path)}")
            msg.attach(part)

        try:
            server = smtplib.SMTP_SSL('smtp.gmail.com', 465)
            server.login(self.email_sender, self.email_password)
            server.send_message(msg)
            server.quit()
            print("✅ Email đã gửi thành công!")
        except Exception as e:
            print(f"❌ Lỗi gửi mail: {e}")

    def cleanup(self):
        if os.path.exists(TEMP_AUDIO_FILE): os.remove(TEMP_AUDIO_FILE)

    def run(self):
        try:
            print("--- BẮT ĐẦU ---")
            ep = self.fetch_latest_episode()       # Bước 1: Lấy thông tin (Đã fix lỗi)
            local = self.download_audio(ep['audio_url']) # Bước 2: Tải file
            ai_content = self.process_with_gemini(local) # Bước 3: AI xử lý
            doc_file = self.create_word_doc(ep, ai_content) # Bước 4: Tạo Word
            self.send_email(doc_file, ep['title']) # Bước 5: Gửi mail
            self.cleanup()
            print("--- THÀNH CÔNG ---")
        except Exception as e:
            print(f"❌ CHƯƠNG TRÌNH THẤT BẠI: {e}")
            self.cleanup()

if __name__ == "__main__":
    PodcastLearningAutomation().run()
